"""
朔州市国控站污染实时通报
csv文件操作
docx文件操作
"""
import os
from datetime import datetime
import operator
from PIL import Image
import cv2
from HTMLTable import (
    HTMLTable,
)
from dateutil.relativedelta import relativedelta
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from sphinx.util import requests
from docx import Document

document = Document()

now = datetime.now()
yes_day2 = datetime.strftime(now - relativedelta(days=1), '%Y%m%d')  # 获取前一天时间
yes_day21 = datetime.strftime(now - relativedelta(days=1), '%Y-%m-%d')  # 获取前一天时间
yes_day22 = datetime.strftime(now - relativedelta(years=1) - relativedelta(days=2), '%Y-%m-%d')  # 获取二天时间
yes_day = datetime.strftime(now - relativedelta(days=1), '%Y-%m-%d 00:00:00')  # 获取前一天时间
today = datetime.now().replace(microsecond=0, minute=0, second=0)  # 当前时间
qntime = (today - relativedelta(years=1))  # 去年当前时间
table = HTMLTable(caption='表1 各属地排名情况')  # 标题
table1 = HTMLTable(caption='表2 各属地排名情况')  # 标题
table2 = '表1 各属地排名情况'
shuozho_url = """
            http://43.254.1.229:8022/AppService/QTService.ashx?method=getcityinfobytype&dateType=day&beginTime={}&endTime={}&userID=xinxiang&AreaID=410701&CityCode=410701&province=&PCityCode=&cityType=all
            """.format(yes_day21, yes_day21)
qn_shuozho_url = """
            http://43.254.1.229:8022/AppService/QTService.ashx?method=getcityinfobytype&dateType=day&beginTime={}&endTime={}&userID=xinxiang&AreaID=410701&CityCode=410701&province=&PCityCode=&cityType=all
            """.format(yes_day22, yes_day22)
png_pm25 = "http://101.201.76.168:8000/picture/pollutant/140/660/daily/PM25_{}.png".format(yes_day2)  # PM25图
png_pm10 = "http://101.201.76.168:8000/picture/pollutant/140/660/daily/PM10_{}.png".format(yes_day2)  # PM10图
imgs = Image.open(requests.get(png_pm25, stream=True).raw)
cap = cv2.VideoCapture(png_pm25)
if (cap.isOpened()):
    ret, img = cap.read()
    cv2.imshow("image", img)
    cv2.waitKey()
result = requests.get(shuozho_url).json()
result1 = requests.get(qn_shuozho_url).json()
#   计算朔州市在山西省的排名
pm25_order_list = []  # PM25 排名
aqi_order_list = []  # AQI 排名
name_list = []  # 市名称
for item in result:
    if item["CityName"] == '太原市':
        name_list.append(item["CityName"])
        PM251 = item['PM25']  # 拿到PM25值
        AQI1 = item['AQI']  # 拿到API值
        pm25_order_list.append(PM251)
        aqi_order_list.append(AQI1)

    elif item["CityName"] == '大同市':
        name_list.append(item["CityName"])
        PM252 = item['PM25']  # 拿到PM25值
        AQI2 = item['AQI']  # 拿到API值

        pm25_order_list.append(PM252)
        aqi_order_list.append(AQI2)

    elif item["CityName"] == '阳泉市':
        name_list.append(item["CityName"])
        PM253 = item['PM25']  # 拿到PM25值
        AQI3 = item['AQI']  # 拿到API值
        pm25_order_list.append(PM253)
        aqi_order_list.append(AQI3)

    elif item["CityName"] == '长治市':
        name_list.append(item["CityName"])
        PM254 = item['PM25']  # 拿到PM25值
        AQI4 = item['AQI']  # 拿到API值
        pm25_order_list.append(PM254)
        aqi_order_list.append(AQI4)

    elif item["CityName"] == '晋城市':
        name_list.append(item["CityName"])
        PM255 = item['PM25']  # 拿到PM25值
        AQI5 = item['AQI']  # 拿到API值
        pm25_order_list.append(PM255)
        aqi_order_list.append(AQI5)

    elif item["CityName"] == '朔州市':
        name_list.append(item["CityName"])
        PM256 = item['PM25']  # 拿到PM25值
        AQI6 = item['AQI']  # 拿到API值
        pm25_order_list.append(PM256)
        aqi_order_list.append(AQI6)

    elif item["CityName"] == '晋中市':
        name_list.append(item["CityName"])
        PM257 = item['PM25']  # 拿到PM25值
        AQI7 = item['AQI']  # 拿到API值
        pm25_order_list.append(PM257)
        aqi_order_list.append(AQI7)

    elif item["CityName"] == '运城市':
        name_list.append(item["CityName"])
        PM258 = item['PM25']  # 拿到PM25值
        AQI8 = item['AQI']  # 拿到API值
        pm25_order_list.append(PM258)
        aqi_order_list.append(AQI8)

    elif item["CityName"] == '忻州市':
        name_list.append(item["CityName"])
        PM259 = item['PM25']  # 拿到PM25值
        AQI9 = item['AQI']  # 拿到API值
        pm25_order_list.append(PM259)
        aqi_order_list.append(AQI9)

    elif item["CityName"] == '临汾市':
        name_list.append(item["CityName"])
        PM2510 = item['PM25']  # 拿到PM25值
        AQI10 = item['AQI']  # 拿到API值
        pm25_order_list.append(PM2510)
        aqi_order_list.append(AQI10)

    elif item["CityName"] == '吕梁市':
        name_list.append(item["CityName"])
        PM2511 = item['PM25']  # 拿到PM25值
        AQI11 = item['AQI']  # 拿到API值
        pm25_order_list.append(PM2511)
        aqi_order_list.append(AQI11)

# pm25_order_list.sort()
# aqi_order_list.sort()
pm25_value = 0  # PM25排名
aqi_vale = 0  # PM10排名
pm25_dict = dict(zip(name_list, pm25_order_list))
pm25_order = sorted(pm25_dict.items(), key=lambda kv: (kv[1], kv[0]))
aqi_dict = dict(zip(name_list, aqi_order_list))
aqi_order = sorted(aqi_dict.items(), key=lambda kv: (kv[1], kv[0]))

for i in range(len(pm25_order)):
    if pm25_order[i][0] == '朔州市':
        pm25_value = i + 1

for j in range(len(aqi_order)):
    if aqi_order[j][0] == '朔州市':
        aqi_vale = j + 1
qn_pm10 = []
qn_pm25 = []
for item in result1:
    if item["CityName"] == '朔州市':
        AQI = int(item['AQI'])  # 拿到API值
        if AQI < 50:
            grade = '优'
        if AQI >= 50 and AQI < 100:
            grade = '良'
        if AQI >= 100 and AQI < 150:
            grade = '轻度污染'
        if AQI >= 150 and AQI < 200:
            grade = '中度污染'
        if AQI >= 200 and AQI < 300:
            grade = '重度污染'
        if AQI > 300:
            grade = '严重污染'

        PM25 = item['PM25']  # 拿到PM25值
        PM10 = item['PM10']  # 拿到PM20值
        qn_pm25.append(PM25)
        qn_pm10.append(PM10)

for item in result:
    if item["CityName"] == '朔州市':
        AQI = int(item['AQI'])  # 拿到API值
        if AQI < 50:
            grade = '优'
        if AQI >= 50 and AQI < 100:
            grade = '良'
        if AQI >= 100 and AQI < 150:
            grade = '轻度污染'
        if AQI >= 150 and AQI < 200:
            grade = '中度污染'
        if AQI >= 200 and AQI < 300:
            grade = '重度污染'
        if AQI > 300:
            grade = '严重污染'
        # 「(当期 / 比较期 -1) * 100%」

        qn_pm25 = ' '.join(qn_pm25)
        qn_pm10 = ' '.join(qn_pm10)

        PM25 = item['PM25']  # 拿到PM25值
        PM10 = item['PM10']  # 拿到PM20值
        pm25_tb = (int(PM25) / int(qn_pm25) - 1) * 100
        pm10_tb = (int(PM10) / int(qn_pm10) - 1) * 100
        pm25_tb = round(pm25_tb, 2)
        pm10_tb = round(pm10_tb, 2)
        pm10_tb = str(pm10_tb)[1:-1]
        pm25_tb = str(pm25_tb)[1:-1]

        data_analysis = '朔州市生态环境局   {}\n昨日浓度情况分析\n昨日我市区AQI为:{}，在山西省内正排第{}名；PM2.5浓度为:{}μg/m3，在山西省内正排第{}名（排名越靠后，PM2.5浓度越高），同比2020年当日浓度下降{}%；PM10浓度为:{}μg/m3，在山西省内正排第7名，同比2020年当日浓度下降{}%；\n10月31日我市整体处于:{}的水平。\n'.format(
            yes_day2, AQI, pm25_value, aqi_vale, PM25, pm10_tb, PM10, pm25_tb, grade)

one_list = []
table.append_header_rows((
    ('乡镇（街道）', 'PM2.5浓度((μg/m3))', '昨日排名'),
))
table1.append_header_rows((
    ('乡镇（街道）', 'PM2.5浓度((μg/m3))', '昨日排名'),
))

# 拿pm25街乡镇数据 pm2.5浓度 和排名
url_pm25 = """
      https://api-adm-data-service.airqualitychina.cn:9998/adm/adm-rank?json={"token":"qiYrxXEv5mJPW6-dU5_wQaCP9QZp3rNp6vobhX5GUfKIuLJ75S1Vka1IfEZNwV8we1Pu5-ZAHa3iPa4oAjQ4g_gOsblH3uvBSNUjW7-dds3HX81g_NeHd2DTDYs82ILCqAserC9Hbje12v_fQmCONVhjBbQJTMu8Be3zreEd3i42yFR_QZymDFGLaU53-VXbxYvKDZEnCoEzoIFxvQpe3jao4-ApmHjvT9k5Y6PFhVY=","is_round":"True","rankid":"125","timeType":"2","time":"%s","auditflag":"0","keyword":""}
      """ % (yes_day)
result = requests.get(url_pm25).json()['result']
street_township_list = []  # 街道
pm25_list = []  # pm25浓度
pm25_rank_list = []  # pm25排名
pm10_list = []  # pm10浓度
pm10_rank_list = []  # pm10排名
for i in result:
    street_township_list.append(i['townname'])
street_township = tuple(street_township_list)

for j in result:
    pm25_rank_list.append((j['vardata'][1]['rank']))  # pm25 排名
    pm25_list.append(int(j['vardata'][1]['value']))  # pm25 浓度
    pm10_rank_list.append(j['vardata'][2]['rank'])  # pm10 排名
    pm10_list.append(int(j['vardata'][2]['value']))  # pm10 浓度

data_tuple = ()
data_list = []
pm_list = []

data_pm10_tuple = ()
data_pm10_list = []
pm10d_list = []

for i in range(len(street_township_list)):  # PM25
    data_tuple = street_township_list[i], str(pm25_list[i]), pm25_rank_list[i]
    data_list.append(list(data_tuple))
    data_pm10_tuple = street_township_list[i], str(pm10_list[i]), pm10_rank_list[i]
    data_pm10_list.append(list(data_pm10_tuple))

data_list.sort(key=operator.itemgetter(2), reverse=False)  # PM25

for j in data_list:  # 去掉 -99 的数
    if j[1] == '-99':
        data_list.remove(j)

data_pm10_list.sort(key=operator.itemgetter(2), reverse=False)  # PM10
# data_pm10_list = list(map(lambda x: (map(lambda a: str(a), x)), data_pm10_list))

for j in data_pm10_list:  # 去掉 -99 的数

    if j[1] == '-99':
        data_pm10_list.remove(j)

for j in range(len(data_list)):  # PM25
    pm_list.append(data_list[j])
    pm10d_list.append(data_pm10_list[j])

# document = Document()


document.add_heading(u'朔州市大气污染热点网格工作日报', 0)

with open('imag_pm25.png', 'wb') as f1:
    image_pm25 = requests.get(png_pm25).content
    f1.write(image_pm25)
with open('imag_pm10.png', 'wb') as f2:
    image_pm10 = requests.get(png_pm10).content
    f2.write(image_pm10)

run_1 = document.add_paragraph().add_run(data_analysis)  # 向文档里添加文字
run_1.font.name = u'仿宋'  # 注：这个好像设置 run 中的西文字体
run_1.font.element.rPr.rFonts.set(qn('w:eastAsia'), u'仿宋')   # 设置中文字体
run_1.font.size = Pt(16)      # 设置字体大小

run_2 = document.add_paragraph().add_run('PM2.5空间分布如下:\n')  # 向文档里添加文字
run_2.font.name = u'仿宋'  # 注：这个好像设置 run 中的西文字体
run_2.font.element.rPr.rFonts.set(qn('w:eastAsia'), u'仿宋')
run_2.font.size = Pt(16)

document.add_picture('imag_pm25.png', width=Inches(3))  # 向文档里添加图片

run_3 = document.add_paragraph().add_run('PM10空间分布如下:\n')  # 向文档里添加文字
run_3.font.name = u'仿宋'  # 注：这个好像设置 run 中的西文字体
run_3.font.element.rPr.rFonts.set(qn('w:eastAsia'), u'仿宋')
run_3.font.size = Pt(16)

document.add_picture('imag_pm10.png', width=Inches(3))  # 向文档里添加图片

document.add_paragraph('表1浓度排名:\n')  # 向文档里添加文字

table = document.add_table(rows=18, cols=3)
hdr_cells = table.rows[0].cells
hdr_cells[0].text = '乡镇（街道）'
hdr_cells[1].text = 'PM2.5浓度((μg/m3))'
hdr_cells[2].text = '昨日排名'

list1 = []
for i in pm_list:
    dict1 = {"name": i}
    list1.append(dict1)

# 17
for i in range(len(list1)):
    for y in range(len(list1[0]["name"])):
        table.cell(i + 1, y).text = str(list1[i]["name"][y])

document.add_paragraph('表2浓度排名:\n')  # 向文档里添加文字
table = document.add_table(rows=18, cols=3)
hdr_cells = table.rows[0].cells
hdr_cells[0].text = '乡镇（街道）'
hdr_cells[1].text = 'PM10浓度((μg/m3))'
hdr_cells[2].text = '昨日排名'
list1 = []
for i in pm10d_list:
    dict1 = {"name": i}
    list1.append(dict1)

#

for i in range(len(list1)):
    for y in range(len(list1[0]["name"])):
        # if i==0:
        #     table.cell(i+1, y).text = str(list1[0]["name"][y])
        # elif i==1:
        #     table.cell(i + 1, y).text = str(list1[1]["name"][y])
        # else:
        table.cell(i + 1, y).text = str(list1[i]["name"][y])

document.save('朔州市大气污染热点网格工作日报.doc')  # 保存文档

os.remove('imag_pm25.png')
os.remove('imag_pm10.png')  # 删除保存在本地的图片
