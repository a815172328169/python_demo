"""
python操作docx文件demo
"""
from docx import Document
from docx.shared import Inches

document = Document()

document.add_heading('Document Title', 0)    # 添加标题

p = document.add_paragraph('A plain paragraph having some ')    # 写入文字
p.add_run('bold').bold = True                 # 追加文字及样式
p.add_run(' and some ')
p.add_run('italic.').italic = True

document.add_heading('Heading, level 1', level=1)    # 添加标题， 1级标题
document.add_paragraph('Intense quote', style='Intense Quote')    # 写入文字 ，style指定样式

document.add_paragraph(
    'first item in unordered list', style='List Bullet'
)             # 写入文字 ，style指定样式=无序列表
document.add_paragraph(
    'first item in ordered list', style='List Number'
)       # 写入文字 ，style指定样式=有序列表

document.add_picture('monty-truth.png', width=Inches(1.25))    # 写入图片

records = (
    (3, '101', 'Spam'),
    (7, '422', 'Eggs'),
    (4, '631', 'Spam, spam, eggs, and spam')
)
# 写入表格
table = document.add_table(rows=1, cols=3)
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Qty'
hdr_cells[1].text = 'Id'
hdr_cells[2].text = 'Desc'
for qty, id, desc in records:
    row_cells = table.add_row().cells
    row_cells[0].text = str(qty)
    row_cells[1].text = id
    row_cells[2].text = desc

document.add_page_break()  # 页码

document.save('demo.docx')